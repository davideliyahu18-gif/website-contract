#!/usr/bin/env python3
"""Generate Word (.docx) and PDF versions of the website service contract."""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

BASE_DIR = Path(__file__).parent
FONT_PATH = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"

SECTIONS = [
    ("1", "מטרת ההסכם", [
        "מטרת הסכם זה היא להסדיר את היחסים המשפטיים והמסחריים בין נותן השירות לבין הלקוח, בקשר עם מתן שירותי תכנון, עיצוב, פיתוח והקמת אתר אינטרנט (\"הפרויקט\") בהתאם לפרטים המפורטים בנספחים ובהסכם זה.",
        "ההסכם מהווה את הבסיס המלא והבלעדי ליחסים בין הצדדים בנוגע לפרויקט, וגובר על כל הצעה, התכתבות, הבנה או הסכמה קודמת, בין בעל פה ובין בכתב.",
    ]),
    ("2", "היקף העבודה", [
        "נותן השירות יבצע עבור הלקוח את השירותים הבאים, בכפוף לתנאי הסכם זה:",
        "• עיצוב גрафי ו-UI/UX – תכנון ממשק משתמש מותאם אישית, כולל עיצוב דפים, צבעוניות, טיפוגרפיה וחוויית משתמש.",
        "• פיתוח ותכנות – בניית האתר, הטמעת תוכן, אינטגרציות נדרשות ופונקציונליות בהתאם למפרט.",
        "• התאמה למובייל (Responsive Design) – התאמת האתר לתצוגה מיטבית במכשירים ניידים, טאבלטים ומחשבים שולחניים.",
        "• SEO בסיסי – אופטימיזציה בסיסית למנועי חיפוש, כולל מטא-תגיות, כותרות, מבנה URL ידידותי ומפת אתר.",
        "• חיבור דומיין ואחסון – במידת הצורך, סיוע בחיבור דומיין, הגדרת DNS והקמת אחסון; עלויות דומיין ואחסון חיצוניות אינן כלולות במחיר, אלא אם נקבע אחרת בכתב.",
        "הערה: כל שירות, תכונה או דרישה שאינה מפורטת במפורש בנספח היקף העבודה או בהסכם זה, תיחשב כעבודה נוספת ותחויב בתשלום נפרד, בהתאם לסעיף 5.",
    ]),
    ("3", "לוחות זמנים ואבני דרך", [
        "3.1. לוח הזמנים לביצוע הפרויקט ייקבע בהתאם לנספח לוחות זמנים המצורף להסכם זה.",
        "3.2. כל עיכוב במסירת חומרים, משוב או אישורים מצד הלקוח, ידחה את לוחות הזמנים בהתאם, ללא אחריות לנותן השירות.",
        "3.3. נותן השירות יעדכן את הלקוח בכתב על התקדמות הפרויקט ובאופן סדיר בכל אבן דרך משמעותית.",
        "שלבים: (1) אפיון וקבלת חומרים, (2) עיצוב ואישור ויז'ואלי, (3) פיתוח ובדיקות, (4) השקה ומסירה סופית.",
    ]),
    ("4", "מחיר ותנאי תשלום", [
        "4.1. סך התמורה עבור הפרויקט: ____________ ₪ (ש\"ח), בתוספת מע\"מ כחוק.",
        "4.2. מבנה תשלומים: מקדמה (30%) עם חתימת ההסכם; תשלום ביניים (40%) עם אישור העיצוב; יתרה (30%) עם מסירת האתר הסופית.",
        "4.3. התשלום יבוצע באמצעות העברה בנקאית / צ'ק / כרטיס אשראי.",
        "4.4. איחור בתשלום מעבר ל-7 ימים – ריבית פיגורים 1.5% לחודש.",
    ]),
    ("5", "שינויים ותוספות עבודה", [
        "5.1. כלול: עד שני סבבי תיקונים לעיצוב, תיקוני באגים בתקופת הבדיקות, והתאמות קלות לתוכן.",
        "5.2. בתשלום נוסף: שינוי מהותי, הוספת דפים, תכונות, אינטגרציות – תעריף שעתי ____________ ₪ או מחיר קבוע מוסכם.",
        "5.3. כל בקשת שינוי תוגש בכתב; הצעת מחיר לפני ביצוע.",
        "5.4. אין חובה לבצע שינויים שטרם אושרו בכתב.",
    ]),
    ("6", "התחייבויות הלקוח", [
        "6.1. מסירת חומרים (תוכן, תמונות, לוגו, מידע עסקי) במועד ובאיכות הנדרשים.",
        "6.2. משוב ואישורים תוך 5 ימי עסקים; אי-מענה ייחשב כאישור.",
        "6.3. החומרים בבעלות הלקוח או ברישיון תקף; ללא הפרת זכויות צד שלישי.",
        "6.4. אחריות הלקוח לנזקים מאיחור או אי-דיוק בחומרים.",
    ]),
    ("7", "זכויות יוצרים וקניין רוחני", [
        "7.1. זכויות היוצרים והקוד יישארו בבעלות נותן השירות עד לתשלום מלא.",
        "7.2. עם תשלום מלא – העברת בעלות מלאה על התוצר, למעט קוד פתוח וספריות צד שלישי.",
        "7.3. נותן השירות רשאי להציג את הפרויקט בפורטפolio, אלא אם נדרש אחרת בכתב.",
        "7.4. במקרה אי-תשלום – נותן השירות רשאי להשתמש, לשכפל או למחוק את התוצר.",
    ]),
    ("8", "אחריות ותמיכה לאחר מסירה", [
        "8.1. תיקון תקלות ובאגים ללא עלות – 30 יום ממועד מסירה (\"תקופת האחריות\").",
        "8.2. לא כולל: שינויי דרישות, עדכוני תוכן, תקלות אחסון/דומיין, פריצות, שימוש לא נכון.",
        "8.3. לאחר תקופת האחריות – חבילת תחזוקה בתנאים נפרדים.",
    ]),
    ("9", "סודיות", [
        "9.1. שמירה על סודיות מידע עסקי, טכני, פיננסי ואישי.",
        "9.2. חובת הסודיות – 3 שנים לפחות לאחר סיום ההסכם.",
        "9.3. לא חלה על מידע שהפך לנחלת הכלל או שנדרש לפי דין.",
    ]),
    ("10", "הגבלת אחריות", [
        "10.1. אחריות מוגבלת לתיקון תקלות בפיתוח; לא כוללת נזקים עקיפים, אובדן רווח או נתונים.",
        "10.2. סך האחריות לא יעלה על סכום התמורה ששולם בפועל.",
        "10.3. לא אחראי לתוכן הלקוח, שירותי צד שלישי, או שימוש בלתי מורשה.",
    ]),
    ("11", "ביטול ההסכם והשלכותיו", [
        "11.1. ביטול בהודעה בכתב של 14 יום מראש.",
        "11.2. ביטול על ידי הלקוח – תשלום עבור עבודה שבוצעה; ללא החזר על מקדמה בגין עבודה שבוצעה.",
        "11.3. ביטול על ידי נותן השירות – במקרה אי-תשלום, הפרה, או חומרים בלתי חוקיים.",
        "11.4. החזרת חומרים ומידע; למעט תשלומים שגובו בדין.",
    ]),
    ("12", "פתרון מחלוקות", [
        "12.1. יישוב מחלוקות במשא ומתן בתום לב.",
        "12.2. לאחר 30 יום – בית משפט מוסמך בישראל; הדין הישראלי יחול.",
        "12.3. חוק החוזים (חלק כללי), תשל\"ג-1973, ודיני מדינת ישראל.",
    ]),
    ("13", "הוראות כלליות", [
        "13.1. שינויים בתוקף רק בכתב וחתימת שני הצדדים.",
        "13.2. אי-מימוש זכות – לא ויתור.",
        "13.3. ההסכם נחתם בשני עותקים מקוריים.",
    ]),
]


def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL direction."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def set_rtl_run(run):
    """Set run to RTL."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)


def add_rtl_paragraph(doc, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6):
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    p.alignment = align
    run = p.add_run(text)
    set_rtl_run(run)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:asciiTheme'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:hAnsiTheme'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
    p.paragraph_format.space_after = Pt(space_after)
    return p


def generate_docx(output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    add_rtl_paragraph(doc, "David Eliyahu Cybersecurity", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_rtl_paragraph(doc, "הסכם למתן שירותי בניית אתר אינטרנט", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Parties
    add_rtl_paragraph(doc, "פרטי נותן השירות", bold=True, size=12, space_after=4)
    provider_fields = [
        "שם העסק: David Eliyahu Cybersecurity",
        "שם מלא: דוד אליהו",
        "ת.ז. / ח.פ.: ___________________________",
        "כתובת: ___________________________",
        "טלפון: ___________________________",
        'דוא"ל: ___________________________',
    ]
    for f in provider_fields:
        add_rtl_paragraph(doc, f, size=10, space_after=2)

    doc.add_paragraph()
    add_rtl_paragraph(doc, "פרטי הלקוח", bold=True, size=12, space_after=4)
    client_fields = [
        "שם העסק: ___________________________",
        "שם מלא: ___________________________",
        "ת.ז. / ח.פ.: ___________________________",
        "כתובת: ___________________________",
        "טלפון: ___________________________",
        'דוא"ל: ___________________________',
    ]
    for f in client_fields:
        add_rtl_paragraph(doc, f, size=10, space_after=2)

    doc.add_paragraph()
    add_rtl_paragraph(doc, "תאריך ההסכם: ___________________________", bold=True, size=11, space_after=12)
    add_rtl_paragraph(doc, "לפיכך הוצהר, הוסכם והותנה בין הצדדים כדלקמן:", size=11, space_after=12)

    for num, title, paragraphs in SECTIONS:
        add_rtl_paragraph(doc, f"{num}. {title}", bold=True, size=12, space_after=4)
        for para in paragraphs:
            add_rtl_paragraph(doc, para, size=10, space_after=4)
        doc.add_paragraph()

    # Signatures
    add_rtl_paragraph(doc, "חתימות הצדדים", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_rtl_paragraph(doc, "ולראיה באו הצדדים על החתום:", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    cells = sig_table.rows[0].cells

    provider_sig = [
        "נותן השירות – David Eliyahu Cybersecurity",
        "",
        "חתימה: ___________________________",
        "שם: דוד אליהו",
        "תאריך: ___________________________",
    ]
    client_sig = [
        "הלקוח",
        "",
        "חתימה: ___________________________",
        "שם: ___________________________",
        "תאריך: ___________________________",
    ]

    for i, lines in enumerate([provider_sig, client_sig]):
        cell = cells[i]
        for line in lines:
            p = cell.add_paragraph(line)
            set_rtl_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                set_rtl_run(run)
                run.font.name = 'Arial'
                run.font.size = Pt(10)

    doc.save(output_path)
    print(f"Created: {output_path}")


def generate_pdf(output_path):
    pdfmetrics.registerFont(TTFont('Hebrew', FONT_PATH))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'HebrewTitle',
        parent=styles['Normal'],
        fontName='Hebrew',
        fontSize=18,
        leading=24,
        alignment=TA_CENTER,
        textColor=HexColor('#0f2744'),
        spaceAfter=12,
    )
    subtitle_style = ParagraphStyle(
        'HebrewSubtitle',
        parent=styles['Normal'],
        fontName='Hebrew',
        fontSize=12,
        leading=16,
        alignment=TA_CENTER,
        textColor=HexColor('#64748b'),
        spaceAfter=20,
    )
    heading_style = ParagraphStyle(
        'HebrewHeading',
        parent=styles['Normal'],
        fontName='Hebrew',
        fontSize=12,
        leading=16,
        alignment=TA_RIGHT,
        textColor=HexColor('#0f2744'),
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        'HebrewBody',
        parent=styles['Normal'],
        fontName='Hebrew',
        fontSize=10,
        leading=15,
        alignment=TA_JUSTIFY,
        spaceAfter=4,
    )
    field_style = ParagraphStyle(
        'HebrewField',
        parent=styles['Normal'],
        fontName='Hebrew',
        fontSize=10,
        leading=14,
        alignment=TA_RIGHT,
        spaceAfter=2,
    )

    story = []

    story.append(Paragraph("David Eliyahu Cybersecurity", subtitle_style))
    story.append(Paragraph("הסכם למתן שירותי בניית אתר אינטרנט", title_style))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("<b>פרטי נותן השירות</b>", heading_style))
    for f in [
        "שם העסק: David Eliyahu Cybersecurity",
        "שם מלא: דוד אליהו",
        "ת.ז. / ח.פ.: ___________________________",
        "כתובת: ___________________________",
        "טלפון: ___________________________",
        'דוא"ל: ___________________________',
    ]:
        story.append(Paragraph(f, field_style))

    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("<b>פרטי הלקוח</b>", heading_style))
    for f in [
        "שם העסק: ___________________________",
        "שם מלא: ___________________________",
        "ת.ז. / ח.פ.: ___________________________",
        "כתובת: ___________________________",
        "טלפון: ___________________________",
        'דוא"ל: ___________________________',
    ]:
        story.append(Paragraph(f, field_style))

    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("<b>תאריך ההסכם:</b> ___________________________", body_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("לפיכך הוצהר, הוסכם והותנה בין הצדדים כדלקמן:", body_style))
    story.append(Spacer(1, 0.3 * cm))

    for num, title, paragraphs in SECTIONS:
        story.append(Paragraph(f"<b>{num}. {title}</b>", heading_style))
        for para in paragraphs:
            story.append(Paragraph(para, body_style))
        story.append(Spacer(1, 0.2 * cm))

    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("<b>חתימות הצדדים</b>", title_style))
    story.append(Paragraph("ולראיה באו הצדדים על החתום:", body_style))
    story.append(Spacer(1, 0.5 * cm))

    sig_data = [
        [
            Paragraph("נותן השירות<br/>David Eliyahu Cybersecurity<br/><br/>חתימה: _______________<br/>שם: דוד אליהו<br/>תאריך: _______________", body_style),
            Paragraph("הלקוח<br/><br/><br/>חתימה: _______________<br/>שם: _______________<br/>תאריך: _______________", body_style),
        ]
    ]
    sig_table = Table(sig_data, colWidths=[8 * cm, 8 * cm])
    sig_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOX', (0, 0), (-1, -1), 0.5, HexColor('#e2e8f0')),
        ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor('#e2e8f0')),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(sig_table)

    doc.build(story)
    print(f"Created: {output_path}")


def generate_pdf_from_html(output_path):
    """Try generating PDF from HTML using puppeteer for better styling."""
    html_path = BASE_DIR / "contract.html"
    script = f"""
const puppeteer = require('puppeteer');
(async () => {{
  const browser = await puppeteer.launch({{ headless: true, args: ['--no-sandbox'] }});
  const page = await browser.newPage();
  await page.goto('file://{html_path.resolve()}', {{ waitUntil: 'networkidle0' }});
  await page.pdf({{
    path: '{output_path.resolve()}',
    format: 'A4',
    printBackground: true,
    margin: {{ top: '15mm', bottom: '15mm', left: '15mm', right: '15mm' }}
  }});
  await browser.close();
  console.log('PDF created from HTML');
}})();
"""
    script_path = BASE_DIR / "_gen_pdf.js"
    script_path.write_text(script, encoding="utf-8")
    import subprocess
    result = subprocess.run(
        ["node", str(script_path)],
        cwd=str(BASE_DIR),
        capture_output=True,
        text=True,
    )
    script_path.unlink(missing_ok=True)
    if result.returncode == 0:
        print(f"Created (HTML-based): {output_path}")
        return True
    print(f"HTML PDF fallback failed: {result.stderr}")
    return False


if __name__ == "__main__":
    docx_path = BASE_DIR / "הסכם-בניית-אתר-David-Eliyahu-Cybersecurity.docx"
    pdf_path = BASE_DIR / "הסכם-בניית-אתר-David-Eliyahu-Cybersecurity.pdf"

    generate_docx(docx_path)

    if not generate_pdf_from_html(pdf_path):
        generate_pdf(pdf_path)

    print("\nAll documents generated successfully.")
